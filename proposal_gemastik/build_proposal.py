from pathlib import Path
import html

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle,
    Image as RLImage, KeepTogether, ListFlowable, ListItem
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "proposal_gemastik"
ASSETS = OUT / "assets"
ASSETS.mkdir(exist_ok=True)

PDF_PATH = OUT / "SoftwareDevelopment_Digiboom.pdf"
DOCX_PATH = OUT / "Proposal_SoftwareDevelopment_Digiboom_JUALIN_AI.docx"
XLSX_PATH = OUT / "Template_Kuesioner_Validasi_JUALIN_AI.xlsx"
NOTES_PATH = OUT / "Catatan_Finalisasi_Submission.txt"
TODAY = "11 Juni 2026"

GREEN = "#16A34A"
BLUE = "#0EA5E9"
PURPLE = "#7C3AED"
ORANGE = "#F97316"
DARK = "#0F172A"
MUTED = "#64748B"
BORDER = "#D9E2EC"
BG = "#F8FAFC"


def font(name="arial.ttf", size=24):
    path = Path(r"C:\Windows\Fonts") / name
    if path.exists():
        return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


F_REG = font("arial.ttf", 22)
F_SMALL = font("arial.ttf", 18)
F_SMALL_BOLD = font("arialbd.ttf", 18)
F_BOLD = font("arialbd.ttf", 24)
F_TITLE = font("arialbd.ttf", 36)
F_SUB = font("arial.ttf", 22)

pdfmetrics.registerFont(TTFont("TNR", r"C:\Windows\Fonts\times.ttf"))
pdfmetrics.registerFont(TTFont("TNR-Bold", r"C:\Windows\Fonts\timesbd.ttf"))
pdfmetrics.registerFont(TTFont("TNR-Italic", r"C:\Windows\Fonts\timesi.ttf"))


def draw_round(draw, xy, r=18, fill="white", outline=BORDER, width=2):
    draw.rounded_rectangle(xy, radius=r, fill=fill, outline=outline, width=width)


def draw_text_wrap(draw, text, xy, max_width, font_obj, fill=DARK, line_gap=6):
    words = text.split()
    lines, cur = [], ""
    for word in words:
        trial = (cur + " " + word).strip()
        if draw.textbbox((0, 0), trial, font=font_obj)[2] <= max_width:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = word
    if cur:
        lines.append(cur)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font_obj, fill=fill)
        y += font_obj.size + line_gap
    return y


def create_architecture():
    img = Image.new("RGB", (1600, 980), BG)
    d = ImageDraw.Draw(img)
    d.text((70, 50), "Arsitektur Multi-Agen JUALIN OS", font=F_TITLE, fill=DARK)
    d.text((72, 96), "Orchestrator merutekan peristiwa ke agen spesialis; setiap aksi tercatat dan dapat disetujui", font=F_SUB, fill=MUTED)
    nodes = [
        (90, 190, 420, 350, "Customer", "Chat publik / WhatsApp / link katalog", BLUE),
        (520, 190, 850, 350, "Orchestrator (Manajer AI)", "Merutekan peristiwa ke agen, kebijakan, laporan harian", PURPLE),
        (950, 190, 1280, 350, "Juru Tawar (Negotiator)", "Mesin nego deterministik, lantai margin, persetujuan", ORANGE),
        (520, 460, 850, 620, "Sales / Inventory / Growth / Finance", "Pramuniaga, Gudang, Marketing, Keuangan, Layanan", GREEN),
        (90, 700, 420, 860, "PostgreSQL + pgvector", "Katalog, order, agent_runs, negotiation_states, memori", "#1D4ED8"),
        (520, 700, 850, 860, "Redis + arq worker", "Cache, rate limit, siklus proaktif terjadwal", "#DC2626"),
        (950, 700, 1280, 860, "Guardrails + HITL + Audit", "Margin floor, anti prompt-injection, approval, audit log", "#0F766E"),
    ]
    for x1, y1, x2, y2, title, desc, color in nodes:
        draw_round(d, (x1, y1, x2, y2), r=24, fill="white", outline="#E2E8F0", width=3)
        d.rounded_rectangle((x1, y1, x1 + 16, y2), radius=8, fill=color)
        d.text((x1 + 36, y1 + 28), title, font=F_BOLD, fill=DARK)
        draw_text_wrap(d, desc, (x1 + 36, y1 + 72), x2 - x1 - 64, F_SMALL, fill=MUTED)
    arrows = [
        ((420, 270), (520, 270)), ((850, 270), (950, 270)), ((1115, 350), (1115, 700)),
        ((950, 540), (850, 540)), ((685, 620), (685, 700)), ((520, 780), (420, 780)),
        ((950, 780), (850, 780)), ((685, 350), (685, 460)), ((950, 270), (820, 460)),
    ]
    for (x1, y1), (x2, y2) in arrows:
        d.line((x1, y1, x2, y2), fill="#94A3B8", width=4)
        if x2 > x1:
            pts = [(x2, y2), (x2 - 14, y2 - 9), (x2 - 14, y2 + 9)]
        elif x2 < x1:
            pts = [(x2, y2), (x2 + 14, y2 - 9), (x2 + 14, y2 + 9)]
        elif y2 > y1:
            pts = [(x2, y2), (x2 - 9, y2 - 14), (x2 + 9, y2 - 14)]
        else:
            pts = [(x2, y2), (x2 - 9, y2 + 14), (x2 + 9, y2 + 14)]
        d.polygon(pts, fill="#94A3B8")
    d.text((70, 915), "Kontrol utama: tenant isolation, rate limiting, idempotency, audit log, dan human takeover.", font=F_SMALL_BOLD, fill=DARK)
    path = ASSETS / "fig01_arsitektur_jualin_ai.png"
    img.save(path, quality=95)
    return path


def create_user_flow():
    img = Image.new("RGB", (1600, 900), BG)
    d = ImageDraw.Draw(img)
    d.text((70, 48), "Alur Penggunaan dan Dampak yang Diukur", font=F_TITLE, fill=DARK)
    d.text((72, 94), "Fokus: seller melihat manfaat dalam 10 menit, customer mendapat respons cepat, order tercatat.", font=F_SUB, fill=MUTED)
    steps = [
        ("1", "Seller daftar", "Input profil toko dan tujuan jualan"),
        ("2", "Input katalog", "Produk, harga, stok, foto, kategori"),
        ("3", "Share link", "Chat publik / WhatsApp / storefront"),
        ("4", "AI melayani", "Jawab produk, rekomendasi, konfirmasi order"),
        ("5", "Order & bayar", "Payment link, follow-up, status order"),
        ("6", "Dashboard", "Revenue, AI assisted order, conversion funnel"),
    ]
    x0, y = 80, 210
    box_w, box_h, gap = 220, 170, 28
    cols = [GREEN, BLUE, PURPLE, ORANGE, "#0F766E", "#334155"]
    for i, (num, title, desc) in enumerate(steps):
        x = x0 + i * (box_w + gap)
        draw_round(d, (x, y, x + box_w, y + box_h), r=22, fill="white", outline="#E2E8F0", width=3)
        d.ellipse((x + 18, y + 18, x + 60, y + 60), fill=cols[i])
        d.text((x + 39, y + 39), num, font=F_SMALL_BOLD, fill="white", anchor="mm")
        d.text((x + 24, y + 82), title, font=F_SMALL_BOLD, fill=DARK)
        draw_text_wrap(d, desc, (x + 24, y + 112), box_w - 48, font("arial.ttf", 17), fill=MUTED)
        if i < len(steps) - 1:
            d.line((x + box_w + 4, y + 85, x + box_w + gap - 6, y + 85), fill="#94A3B8", width=4)
            d.polygon([(x + box_w + gap - 6, y + 85), (x + box_w + gap - 20, y + 76), (x + box_w + gap - 20, y + 94)], fill="#94A3B8")
    d.text((90, 500), "Metrik evaluasi", font=F_BOLD, fill=DARK)
    metrics = [
        ("Respons", "Waktu balas rata-rata, target detik"),
        ("Konversi", "Chat yang berujung order"),
        ("Kepercayaan", "Kebijakan toko dan payment jelas"),
        ("Kontrol", "Manual takeover dan approval action"),
    ]
    for i, (m, desc) in enumerate(metrics):
        x = 90 + i * 360
        draw_round(d, (x, 555, x + 310, 720), r=18, fill="#FFFFFF", outline="#DCE6F2", width=2)
        d.text((x + 22, 580), m, font=F_BOLD, fill=GREEN)
        draw_text_wrap(d, desc, (x + 22, 626), 260, F_SMALL, fill=MUTED)
    d.text((90, 792), "Validasi primer: kuesioner 30 responden asli dari GForm, dilengkapi uji skenario prototipe.", font=F_SMALL_BOLD, fill=DARK)
    path = ASSETS / "fig02_alur_pengguna_dampak.png"
    img.save(path, quality=95)
    return path


def bubble_lines(draw, text, maxw, fnt):
    lines, cur = [], ""
    for w in text.split():
        trial = (cur + " " + w).strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= maxw:
            cur = trial
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


def create_chat_mockup():
    img = Image.new("RGB", (1500, 980), "#ECFDF5")
    d = ImageDraw.Draw(img)
    d.text((70, 50), "Mockup Interface: Customer Chat", font=F_TITLE, fill=DARK)
    d.text((72, 96), "Percakapan dibuat ringkas, katalog-aware, dan selalu mengarah ke konfirmasi order.", font=F_SUB, fill=MUTED)
    px, py, pw, ph = 160, 170, 470, 710
    draw_round(d, (px, py, px + pw, py + ph), r=42, fill="white", outline="#BBF7D0", width=4)
    d.rounded_rectangle((px + 20, py + 20, px + pw - 20, py + 88), radius=24, fill="#DCFCE7")
    d.ellipse((px + 42, py + 38, px + 76, py + 72), fill=GREEN)
    d.text((px + 92, py + 37), "Toko Sari Fashion", font=F_SMALL_BOLD, fill=DARK)
    d.text((px + 92, py + 62), "Online - dibantu JUALIN.AI", font=font("arial.ttf", 14), fill=MUTED)
    fnt = font("arial.ttf", 18)
    bubbles = [
        ("left", "Kak, ada baju buat kondangan warna soft?"),
        ("right", "Ada kak. Rekomendasi paling cocok: Baju Pink Satin Rp 89.000, stok M dan L ready. Bahannya ringan dan rapi untuk acara."),
        ("left", "Ukuran M ready? mau order 1."),
        ("right", "Ready kak. Saya konfirmasi dulu: Baju Pink Satin ukuran M, qty 1, total Rp 89.000. Mau lanjut dibuatkan pesanan?"),
    ]
    y = py + 120
    for side, text in bubbles:
        maxw = 330
        lines = bubble_lines(d, text, maxw, fnt)
        h = 32 + len(lines) * 25
        x = px + 36 if side == "left" else px + pw - 36 - maxw - 24
        fill = "#F1F5F9" if side == "left" else "#DCFCE7"
        outline = "#E2E8F0" if side == "left" else "#86EFAC"
        draw_round(d, (x, y, x + maxw + 24, y + h), r=22, fill=fill, outline=outline, width=2)
        ty = y + 17
        for line in lines:
            d.text((x + 14, ty), line, font=fnt, fill=DARK)
            ty += 25
        y += h + 18
    d.rounded_rectangle((px + 28, py + ph - 72, px + pw - 28, py + ph - 28), radius=22, fill="#F8FAFC", outline="#E2E8F0")
    d.text((px + 52, py + ph - 60), "Tulis pesan...", font=font("arial.ttf", 17), fill="#94A3B8")
    d.ellipse((px + pw - 76, py + ph - 67, px + pw - 36, py + ph - 27), fill=GREEN)
    d.polygon([(px + pw - 60, py + ph - 57), (px + pw - 46, py + ph - 47), (px + pw - 60, py + ph - 37)], fill="white")
    x = 720
    d.text((x, 180), "Elemen interface yang ditonjolkan", font=F_BOLD, fill=DARK)
    cards = [
        ("Jawaban berbasis katalog", "AI tidak mengarang harga/stok; semua merujuk produk seller."),
        ("Konfirmasi sebelum order", "Pesanan baru dibuat setelah customer menyetujui detail item."),
        ("Tone lokal dan natural", "Bahasa singkat, sopan, cocok untuk chat jual-beli UMKM."),
        ("Fallback aman", "Jika LLM gagal atau data tidak lengkap, sistem meminta admin mengecek manual."),
    ]
    yy = 240
    for title, desc in cards:
        draw_round(d, (x, yy, 1370, yy + 130), r=20, fill="white", outline="#BBF7D0", width=2)
        d.text((x + 28, yy + 24), title, font=F_SMALL_BOLD, fill=GREEN)
        draw_text_wrap(d, desc, (x + 28, yy + 58), 570, F_SMALL, fill=MUTED)
        yy += 155
    path = ASSETS / "mockup01_customer_chat.png"
    img.save(path, quality=95)
    return path


def create_dashboard_mockup():
    img = Image.new("RGB", (1600, 1000), BG)
    d = ImageDraw.Draw(img)
    d.text((70, 50), "Mockup Interface: Seller Dashboard", font=F_TITLE, fill=DARK)
    d.text((72, 96), "Dashboard fokus pada uang, order, stok, dan chat yang perlu tindakan.", font=F_SUB, fill=MUTED)
    sx, sy, sw, sh = 80, 170, 1440, 730
    draw_round(d, (sx, sy, sx + sw, sy + sh), r=28, fill="white", outline="#E2E8F0", width=3)
    d.rounded_rectangle((sx, sy, sx + 250, sy + sh), radius=28, fill="#0F172A")
    d.text((sx + 32, sy + 34), "JUALIN.AI", font=F_BOLD, fill="white")
    nav = ["Overview", "Inbox", "Produk", "Order", "Analytics", "Campaign", "Settings"]
    for i, n in enumerate(nav):
        y = sy + 100 + i * 54
        fill = GREEN if i == 0 else "#1E293B"
        d.rounded_rectangle((sx + 24, y, sx + 226, y + 40), radius=12, fill=fill)
        d.text((sx + 48, y + 10), n, font=font("arial.ttf", 17), fill="white")
    cx = sx + 290
    d.text((cx, sy + 38), "Selamat datang, Seller", font=F_BOLD, fill=DARK)
    d.text((cx, sy + 72), "Ringkasan performa toko hari ini", font=F_SMALL, fill=MUTED)
    stats = [
        ("AI Bantu Closing", "Rp 1,8 Jt", GREEN),
        ("Order Dibantu AI", "23", BLUE),
        ("Payment Pending", "Rp 420 rb", ORANGE),
        ("Avg Respons", "3 dtk", PURPLE),
    ]
    for i, (label, val, color) in enumerate(stats):
        x = cx + i * 285
        draw_round(d, (x, sy + 125, x + 255, sy + 250), r=18, fill=BG, outline="#E2E8F0", width=2)
        d.rectangle((x, sy + 125, x + 8, sy + 250), fill=color)
        d.text((x + 26, sy + 150), label, font=font("arial.ttf", 16), fill=MUTED)
        d.text((x + 26, sy + 182), val, font=font("arialbd.ttf", 28), fill=DARK)
    draw_round(d, (cx, sy + 300, cx + 720, sy + 650), r=18, fill="white", outline="#E2E8F0", width=2)
    d.text((cx + 28, sy + 326), "Order 7 Hari Terakhir", font=F_SMALL_BOLD, fill=DARK)
    bars = [90, 130, 80, 180, 220, 160, 250]
    for i, b in enumerate(bars):
        x = cx + 70 + i * 82
        d.rounded_rectangle((x, sy + 610 - b, x + 42, sy + 610), radius=12, fill=GREEN)
        d.text((x + 8, sy + 622), ["Min", "Sen", "Sel", "Rab", "Kam", "Jum", "Sab"][i], font=font("arial.ttf", 14), fill=MUTED)
    draw_round(d, (cx + 760, sy + 300, sx + sw - 40, sy + 650), r=18, fill="#F0FDF4", outline="#BBF7D0", width=2)
    d.text((cx + 790, sy + 328), "Tindakan Prioritas", font=F_SMALL_BOLD, fill=DARK)
    tasks = ["Follow-up 8 order pending", "Upload foto 5 produk", "Cek 3 chat perlu admin", "Publish storefront minggu ini"]
    yy = sy + 380
    for task in tasks:
        d.rounded_rectangle((cx + 790, yy, sx + sw - 80, yy + 50), radius=12, fill="white", outline="#BBF7D0")
        d.ellipse((cx + 812, yy + 16, cx + 830, yy + 34), fill=GREEN)
        d.text((cx + 846, yy + 14), task, font=font("arial.ttf", 17), fill=DARK)
        yy += 66
    path = ASSETS / "mockup02_seller_dashboard.png"
    img.save(path, quality=95)
    return path


def create_onboarding_mockup():
    img = Image.new("RGB", (1500, 900), BG)
    d = ImageDraw.Draw(img)
    d.text((70, 50), "Mockup Interface: Quick Start 10 Menit", font=F_TITLE, fill=DARK)
    d.text((72, 96), "Onboarding mobile-first agar UMKM tidak perlu memahami konfigurasi teknis.", font=F_SUB, fill=MUTED)
    steps = ["Toko", "Produk", "AI", "Preview", "Go Live"]
    x0, y0 = 110, 190
    for i, step in enumerate(steps):
        x = x0 + i * 260
        color = GREEN if i <= 2 else "#CBD5E1"
        d.ellipse((x, y0, x + 54, y0 + 54), fill=color)
        d.text((x + 27, y0 + 27), str(i + 1), font=F_SMALL_BOLD, fill="white", anchor="mm")
        d.text((x - 8, y0 + 70), step, font=F_SMALL_BOLD, fill=DARK)
        if i < len(steps) - 1:
            d.line((x + 64, y0 + 27, x + 250, y0 + 27), fill="#CBD5E1", width=4)
    cards = [
        ("Profil Toko", "Nama toko, kategori usaha, jam layanan, gaya bahasa."),
        ("Produk Cepat", "Tambah 3-5 produk pertama dari HP, bisa draft dulu."),
        ("Simulasi Chat", "Seller mengetes pertanyaan customer sebelum link dibagikan."),
    ]
    for i, (title, desc) in enumerate(cards):
        x = 130 + i * 430
        draw_round(d, (x, 360, x + 360, 665), r=26, fill="white", outline="#E2E8F0", width=3)
        d.rounded_rectangle((x + 26, 390, x + 110, 474), radius=20, fill="#DCFCE7")
        d.text((x + 52, 414), str(i + 1), font=F_TITLE, fill=GREEN)
        d.text((x + 26, 510), title, font=F_BOLD, fill=DARK)
        draw_text_wrap(d, desc, (x + 26, 552), 305, F_SMALL, fill=MUTED)
    d.rounded_rectangle((130, 735, 1370, 805), radius=22, fill="#0F172A")
    d.text((170, 756), "Output onboarding: link katalog, halaman chat demo, dan checklist siap jual.", font=F_SMALL_BOLD, fill="white")
    path = ASSETS / "mockup03_onboarding_quick_start.png"
    img.save(path, quality=95)
    return path


asset_paths = [
    create_architecture(),
    create_user_flow(),
    create_chat_mockup(),
    create_dashboard_mockup(),
    create_onboarding_mockup(),
]

refs = [
    "Yao, S., Zhao, J., Yu, D., Du, N., Shafran, I., Narasimhan, K., & Cao, Y. (2023). ReAct: Synergizing Reasoning and Acting in Language Models. International Conference on Learning Representations (ICLR). https://arxiv.org/abs/2210.03629",
    "Shinn, N., Cassano, F., Gopinath, A., Narasimhan, K., & Yao, S. (2023). Reflexion: Language Agents with Verbal Reinforcement Learning. Advances in Neural Information Processing Systems (NeurIPS). https://arxiv.org/abs/2303.11366",
    "Xia, T., Yu, B., Dang, K., et al. (2024). Measuring Bargaining Abilities of Large Language Models: A Benchmark and a Buyer-Enhancement Method. Findings of the Association for Computational Linguistics (ACL). https://arxiv.org/abs/2402.15813",
    "Bianchi, F., Chia, P. J., Yuksekgonul, M., et al. (2024). How Well Can LLMs Negotiate? NegotiationArena Platform and Analysis. https://arxiv.org/abs/2402.05863",
    "Wang, L., Ma, C., Feng, X., et al. (2024). A Survey on Large Language Model based Autonomous Agents. Frontiers of Computer Science, 18(6). https://doi.org/10.1007/s11704-024-40231-1",
    "Lewis, P., Perez, E., Piktus, A., et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS. https://arxiv.org/abs/2005.11401",
    "Badan Pusat Statistik. (2024). Statistik E-Commerce 2024. Jakarta: BPS. https://www.bps.go.id/",
    "Google, Temasek, & Bain & Company. (2024). e-Conomy SEA 2024. https://economysea.withgoogle.com/",
    "Kementerian Koperasi dan UKM Republik Indonesia. (2023). Perkembangan Data Usaha Mikro, Kecil, dan Menengah. https://kemenkopukm.go.id/",
    "Bank Indonesia. (2024). Laporan Perkembangan Ekonomi dan Keuangan Digital serta QRIS. https://www.bi.go.id/",
    "NIST. (2023). Artificial Intelligence Risk Management Framework (AI RMF 1.0). https://www.nist.gov/itl/ai-risk-management-framework",
    "International Organization for Standardization. (2023). ISO/IEC 25010 software quality model. https://www.iso.org/standard/78176.html",
    "pgvector Contributors. (2024). pgvector: Open-source vector similarity search for PostgreSQL. https://github.com/pgvector/pgvector",
]

sections = []


def add_section(title, blocks):
    sections.append((title, blocks))


add_section("Ringkasan", [
    ("p", "Kami memulai JUALIN.AI dari hal sederhana yang kami lihat sendiri. Teman-teman kami yang berjualan online lewat WhatsApp selalu kewalahan. Chat masuk tengah malam, pembeli menawar, stok lupa diperbarui, dan tidak sedikit orderan batal hanya karena telat dibalas. Dari situ muncul pertanyaan kami, bagaimana kalau penjual kecil punya karyawan yang tidak pernah tidur? JUALIN.AI dalam edisi JUALIN OS adalah jawaban kami, yaitu bukan sekadar chatbot yang menjawab, melainkan satu tim agen AI yang menjalankan toko: melayani percakapan, menawar harga secara aman tanpa membuat penjual rugi, menjaga stok, menagih pembayaran, dan menyusun laporan harian, sementara keputusan penting tetap dipegang penjual. Proposal ini memaparkan latar masalah, rancangan, dan prototipe yang sudah berjalan, beserta rencana pengukuran dampaknya bagi UMKM mikro Indonesia."),
])

add_section("1. Judul/Nama Perangkat Lunak", [
    ("p", "Nama perangkat lunak yang diusulkan adalah JUALIN.AI dengan edisi JUALIN OS, sebuah Sistem Operasi Toko Otonom untuk UMKM mikro. Berbeda dari chatbot layanan yang hanya menjawab chat, JUALIN OS adalah tim agen kecerdasan buatan terspesialisasi yang menjalankan operasional toko secara otonom dan terkoordinasi, yaitu melayani percakapan, menawar harga secara aman, menjaga stok, menagih pembayaran, dan menyusun laporan harian, dengan kendali penuh tetap berada di tangan penjual. Produk ini dikembangkan untuk Divisi III Pengembangan Perangkat Lunak GEMASTIK DIGINEXS 2026 dengan tema Digital Intelligence For Smart Society."),
    ("table", [
        ["Aspek", "Rincian"],
        ["Nama produk", "JUALIN.AI - edisi JUALIN OS (Sistem Operasi Toko Otonom)"],
        ["Tagline", "Tim karyawan AI yang menjalankan toko UMKM: menjual, menawar, menjaga stok, menagih, dan membukukan secara otonom dan terkendali"],
        ["Target pengguna", "Pelaku UMKM mikro yang berjualan melalui WhatsApp, Instagram, storefront, atau marketplace sosial"],
        ["Kebaruan inti", "Arsitektur multi-agen (Manajer dan enam agen spesialis) serta mesin negosiasi aman-margin yang dijamin tidak pernah menjual di bawah batas untung"],
        ["Kontribusi SDGs", "SDG 8 Decent Work and Economic Growth, SDG 9 Industry Innovation and Infrastructure, SDG 1 No Poverty, SDG 10 Reduced Inequality"],
        ["Status prototipe", "Aplikasi web full-stack operasional: Next.js, FastAPI, PostgreSQL pgvector, Redis, worker arq, Docker, dengan modul multi-agen dan dashboard AI Crew"],
    ], [3.6, 9.6]),
    ("p", "Inti gagasan JUALIN OS adalah memindahkan UMKM dari posisi AI Copilot yang sekadar membantu menjawab menjadi AI Autopilot yang menjalankan operasi. Sebuah Orchestrator merutekan setiap peristiwa, mesin negosiasi deterministik menjaga margin, dan setiap tindakan agen tercatat serta dapat disetujui atau ditolak penjual. Dengan demikian otomasi menjadi lebih berdampak sekaligus tetap dapat dipercaya."),
    ("table", [
        ["Kriteria penyisihan", "Respons dalam proposal"],
        ["Aspek inovasi 20%", "Lompatan kategori dari chatbot reaktif menjadi sistem multi-agen otonom dengan negosiasi aman-margin; berakar pada riset ReAct, Reflexion, dan bargaining LLM."],
        ["Dampak 20%", "Dampak terukur: waktu operasional dihemat, omzet dibantu AI, omzet diselamatkan, dan deal negosiasi tertutup tanpa jual-rugi; divalidasi melalui kuesioner 30 responden asli."],
        ["Desain antarmuka 20%", "Pusat Komando AI Crew menampilkan status tiap agen, activity feed, antrean persetujuan, dan laporan harian; mobile-first untuk seller dan chat untuk customer."],
        ["Proses pengembangan 20%", "Agile prototyping, arsitektur multi-agen modular, guardrail deterministik, savepoint transaksi, feature flag, dan CI/CD GitHub Actions."],
        ["Kesesuaian ide 10%", "Selaras tema smart society: memberdayakan jutaan UMKM dengan kecerdasan digital yang inklusif."],
        ["Urgensi masalah 10%", "Penjual mikro menanggung enam pekerjaan sekaligus pada kanal chat yang tidak pernah berhenti."],
    ], [3.5, 9.7]),
])

add_section("2. Latar Belakang Ide Perangkat Lunak", [
    ("p", "Ide ini tidak kami dapat dari laporan riset, melainkan dari apa yang kami lihat sehari-hari. Banyak teman dan tetangga kami berjualan baju, makanan, dan barang harian lewat WhatsApp dan Instagram. Mereka pandai membuat produk, tetapi kewalahan mengurus chat. Pesan menumpuk saat mereka sedang masak, mengantar anak, atau tidur. Saat ditawar, mereka ragu boleh turun berapa supaya tidak rugi. Orderan yang sudah dibuat pun sering batal hanya karena lupa ditagih. Dari situ kami sadar, masalahnya bukan pada produk, tetapi pada keterbatasan satu orang melayani semuanya sekaligus."),
    ("p", "Pengamatan kami ternyata sejalan dengan data. UMKM menyumbang sekitar 60 persen Produk Domestik Bruto dan menyerap hampir 97 persen tenaga kerja Indonesia menurut Kementerian Koperasi dan UKM. Dari sekitar 4,4 juta pelaku e-commerce nasional, mayoritas justru berjualan lewat pesan instan, bukan etalase marketplace, menurut Badan Pusat Statistik tahun 2024. Artinya jutaan transaksi UMKM benar-benar terjadi di dalam kotak chat, dan di situlah penjual paling mudah kehabisan tenaga."),
    ("p", "Kalau dirinci, seorang penjual mikro sebenarnya mengerjakan enam peran sekaligus, sendirian:"),
    ("table", [
        ["Beban penjual (enam pekerjaan)", "Dampak bila tak tertangani", "Agen JUALIN OS yang menangani"],
        ["Balas chat selama 24 jam", "Telat balas, pembeli pindah ke toko lain", "Pramuniaga (Sales) berbasis katalog"],
        ["Tawar-menawar harga", "Salah hitung, terlanjur jual rugi", "Juru Tawar (Negotiator) aman-margin"],
        ["Tagih pembayaran tertunda", "Omzet menguap", "Marketing (Growth) proaktif"],
        ["Jaga stok dan cegah oversell", "Janji palsu, komplain pelanggan", "Gudang (Inventory)"],
        ["Pembukuan harian", "Tidak tahu performa toko", "Keuangan (Finance) dan Laporan Harian"],
        ["Kendali dan kepercayaan", "Takut AI salah ambil keputusan", "Manajer (Orchestrator) dan persetujuan manusia"],
    ], [3.4, 4.6, 5.2]),
    ("p", "Kami sempat mencoba beberapa chatbot yang sudah ada. Hampir semuanya reaktif, yaitu hanya menunggu pertanyaan lalu menjawab. Tidak ada yang berinisiatif menagih, menyapa pelanggan lama, apalagi berani menawar harga, padahal menawar justru bagian paling khas dari jual-beli di Indonesia. Dari sinilah kami yakin ada celah yang belum diisi siapa pun, yaitu AI yang benar-benar menjalankan toko di sisi penjual, termasuk menawar dengan aman."),
    ("p", "Arah ini juga didukung tren global. Perdagangan yang dijalankan agen AI diperkirakan menaikkan konversi sekitar 20 sampai 30 persen dan menjadi pasar di atas 50 miliar dolar Amerika Serikat pada tahun 2030, tetapi hampir semuanya masih berfokus membantu pembeli, bukan penjual kecil. Kami merancang JUALIN OS untuk mengisi ruang kosong itu, dengan konteks dan bahasa yang akrab bagi UMKM Indonesia."),
])

add_section("3. Tujuan dan Manfaat Dikembangkan Perangkat Lunak", [
    ("p", "Tujuan utama JUALIN OS adalah memberi UMKM mikro sebuah tim agen AI yang menjalankan operasional toko secara otonom namun tetap dikendalikan penjual, sehingga penjual dapat melayani lebih banyak pembeli, menutup lebih banyak order, dan tidak kehilangan omzet karena keterbatasan waktu."),
    ("bullets", [
        "Menyediakan agen Pramuniaga yang menjawab pertanyaan customer berdasarkan katalog, harga, stok, dan kebijakan toko.",
        "Menyediakan agen Juru Tawar yang menanggapi tawar-menawar dengan penawaran yang dijamin tidak pernah di bawah batas margin penjual.",
        "Menjalankan agen Marketing dan Keuangan secara proaktif untuk menagih pembayaran tertunda, menyapa pelanggan pasif, dan menyusun laporan harian.",
        "Menjaga kepercayaan melalui human-in-the-loop, activity feed yang transparan, audit log, dan kendali kebijakan per penjual.",
        "Mengukur dampak nyata seperti waktu operasional dihemat, omzet dibantu AI, omzet diselamatkan, dan konversi chat menjadi order.",
    ]),
    ("table", [
        ["Pihak", "Manfaat"],
        ["UMKM/seller", "Beban enam pekerjaan terbantu otomatis, respons instan, negosiasi aman, penagihan tidak terlewat, dan laporan harian otomatis, tanpa menambah karyawan."],
        ["Customer", "Dilayani cepat selama 24 jam, mendapat penawaran tawar-menawar yang wajar, dan proses order yang jelas."],
        ["Perguruan tinggi dan ekosistem", "Bukti karya software bernilai tinggi yang menerapkan AI agentik secara bertanggung jawab pada masalah ekonomi digital lokal."],
        ["Masyarakat (SDGs)", "Mendorong pertumbuhan ekonomi inklusif (SDG 8), inovasi infrastruktur digital (SDG 9), dan pengurangan ketimpangan akses teknologi (SDG 1 dan SDG 10)."],
    ], [3.0, 10.2]),
    ("p", "Dampak ditargetkan terukur, bukan sekadar argumentasi. Indikator keberhasilan awal mencakup penghematan waktu operasional penjual, peningkatan konversi chat menjadi order, nilai pembayaran tertunda yang berhasil diselamatkan, serta jumlah deal negosiasi yang tertutup tanpa satu pun transaksi di bawah batas margin."),
])

add_section("4. Batasan Perangkat Lunak yang Dikembangkan", [
    ("p", "Batasan dibuat agar prototipe tetap realistis, aman, dan dapat diuji dalam waktu kompetisi. JUALIN.AI tidak diposisikan sebagai AI yang menggantikan seluruh keputusan bisnis seller, tetapi sebagai asisten operasional yang mengurangi beban chat dan administrasi order."),
    ("bullets", [
        "AI hanya boleh menjawab berdasarkan data katalog, kebijakan, dan konfigurasi seller yang tersedia di sistem.",
        "Aksi berdampak uang seperti diskon, broadcast campaign, dan perubahan order sensitif memerlukan approval atau konfigurasi eksplisit dari seller.",
        "Integrasi WhatsApp Cloud, Midtrans, dan Cashi.id disiapkan sebagai plugin/konfigurasi; jika credential belum tersedia, sistem tetap berjalan melalui public chat demo.",
        "Prototipe tidak menangani logistik kompleks lintas ekspedisi; v1 menggunakan shipping cost manual dan kebijakan toko.",
        "Sistem tidak memberi nasihat hukum, medis, finansial, atau topik di luar transaksi jual-beli toko.",
        "Validasi responden harus menggunakan jawaban asli dari GForm; data simulasi hanya boleh dipakai untuk uji format rekap, bukan sebagai bukti penelitian.",
    ]),
])

add_section("5. Metodologi Pengembangan Perangkat Lunak", [
    ("p", "Metodologi yang digunakan adalah agile prototyping dengan tahap discovery, design, build, test, dan validate. Pendekatan ini dipilih karena produk perlu cepat diuji pada skenario chat penjualan nyata, namun tetap menuntut kontrol kualitas tinggi pada keamanan, guardrail negosiasi, dan isolasi multi-tenant. Fitur berisiko dibangun di belakang feature flag dan diintegrasikan bertahap sebagai evolusi, bukan penulisan ulang."),
    ("table", [
        ["Tahap", "Aktivitas", "Output"],
        ["Discovery", "Memetakan enam beban penjual UMKM serta meriset referensi agen AI dan negosiasi model bahasa.", "Problem statement, persona, dan metrik dampak."],
        ["Design", "Merancang arsitektur multi-agen, mesin negosiasi deterministik, dan antarmuka AI Crew.", "Diagram arsitektur, desain solusi, dan mockup."],
        ["Build", "Mengembangkan modul services agent_os, tabel agent_os, API agent-os, dashboard AI Crew, dan worker proaktif di atas basis kode existing.", "Prototipe multi-agen yang dapat dijalankan."],
        ["Test", "Menjalankan compile check, validasi mesin negosiasi, serta pengujian guardrail dan keamanan melalui CI.", "Bukti verifikasi teknis."],
        ["Validate", "Mengumpulkan respons 30 responden asli melalui GForm dan menguji demo pada calon penjual atau pembeli.", "Rekap minat, keberatan, dan prioritas fitur."],
    ], [2.2, 6.1, 4.9]),
    ("p", "Verifikasi teknis yang telah dilakukan mencakup kompilasi seluruh modul backend melalui compileall dan validasi matematis mesin negosiasi yang menunjukkan nol pelanggaran lantai margin pada seluruh skenario dan ronde uji. Pipeline CI GitHub Actions dikonfigurasi menjalankan pytest, audit dependensi, lint, dan build frontend pada setiap perubahan, sehingga kualitas terjaga otomatis sebelum deploy."),
])

add_section("6. Analisis Kebutuhan dan Desain Solusi Perangkat Lunak", [
    ("p", "Aktor utama sistem adalah seller, customer, dan admin. Yang membedakan JUALIN OS adalah lapisan agen otonom di antara mereka, yaitu sebuah Orchestrator yang merutekan setiap peristiwa seperti chat masuk, pembayaran, perubahan stok, dan jadwal terjadwal ke agen spesialis yang tepat, lalu mencatat tindakannya untuk diaudit dan, bila berisiko, dimintakan persetujuan penjual."),
    ("table", [
        ["Agen (peran)", "Tanggung jawab otonom"],
        ["Orchestrator (Manajer)", "Merutekan peristiwa ke agen, menjaga kebijakan global, dan menyusun Laporan Harian."],
        ["Sales (Pramuniaga)", "Melayani percakapan berbasis katalog: menyapa, menggali kebutuhan, presentasi, dan closing."],
        ["Negotiator (Juru Tawar)", "Menanggapi tawar-menawar dengan penawaran yang dijamin tidak pernah di bawah lantai margin."],
        ["Inventory (Gudang)", "Memverifikasi stok sebelum janji atau order, dan mendeteksi stok menipis."],
        ["Growth (Marketing)", "Menagih pembayaran tertunda dan menyapa kembali pelanggan pasif secara proaktif."],
        ["Finance (Keuangan)", "Merekap omzet, pembayaran lunas dibandingkan tertunda, dan produk terlaris."],
    ], [3.4, 9.8]),
    ("figure", str(asset_paths[0]), "Gambar 1. Arsitektur multi-agen JUALIN OS."),
    ("figure", str(asset_paths[1]), "Gambar 2. Alur penggunaan dan dampak yang diukur."),
    ("p", "Kebaruan teknis terpenting adalah pemisahan angka dari kata pada negosiasi. Sebuah penghasil penawaran yang deterministik menghitung tawaran balik memakai tangga konsesi yang dibatasi lantai harga, yaitu nilai terbesar antara modal dikali satu tambah margin minimum dan harga dikali satu kurang diskon maksimum. Model bahasa hanya merangkai kalimat di sekitar angka yang sudah diputuskan mesin, sehingga AI tidak mungkin mengarang harga atau menjual rugi. Pendekatan ini sejalan dengan temuan riset bargaining model bahasa bahwa agen membutuhkan penghasil penawaran untuk mengontrol rentang harga."),
    ("table", [
        ["Modul", "Desain solusi"],
        ["Mesin negosiasi", "Mesin deterministik menjaga lantai margin; model bahasa hanya untuk bahasa; diskon di atas ambang masuk antrean persetujuan."],
        ["Memori dan state", "Memori customer berbasis pgvector untuk personalisasi; negotiation_states menyimpan konteks tawar yang sedang berjalan."],
        ["Substrat peristiwa", "background_jobs dan worker arq untuk kerja proaktif; audit_logs dan agent_runs untuk jejak yang dapat diaudit."],
        ["Kendali manusia", "agent_policies sebagai sakelar per agen dan ambang diskon; agent_approvals sebagai antrean human-in-the-loop."],
        ["Ketahanan", "Setiap pemanggilan agen dibungkus savepoint transaksi sehingga kegagalan agen tidak menjatuhkan chat utama."],
        ["Observability", "AI trace, usage event, dan provider health untuk pengawasan operasional."],
    ], [3.0, 10.2]),
])

add_section("7. Implementasi Perangkat Lunak", [
    ("p", "Implementasi JUALIN OS berbentuk aplikasi web full-stack yang dapat dijalankan, dibangun secara bertahap di atas basis kode JUALIN.AI. Lapisan multi-agen ditambahkan melalui modul models agent_os dengan empat tabel yaitu AgentPolicy, AgentRun, AgentApproval, dan NegotiationState, paket services agent_os yang berisi orchestrator, negotiation, inventory, finance, growth, brief, dan cycles, API agent-os, worker terjadwal, serta halaman dashboard AI Crew. Seluruh fitur baru dikendalikan feature flag sehingga aman terhadap perilaku lama."),
    ("table", [
        ["Layer", "Teknologi", "Alasan pemilihan"],
        ["Frontend", "Next.js 16, React, CSS Modules", "Cepat untuk dashboard dan chat publik, mendukung UI mobile-first."],
        ["Backend", "FastAPI, Pydantic, SQLAlchemy async", "Ringan, cocok untuk API modular dan orkestrasi agen, validasi request jelas."],
        ["Database", "PostgreSQL 16 + pgvector", "Data relasional kuat, sementara pencarian semantik katalog dilakukan pada basis data yang sama."],
        ["Cache dan worker", "Redis 7 dan arq", "Rate limit, cache katalog, dan kerja proaktif agen yang terjadwal dan hemat."],
        ["AI", "Model bahasa via endpoint OpenAI-compatible, embedding all-MiniLM-L6-v2, guardrails dan mesin deterministik", "Respons natural dengan angka harga dikontrol mesin, biaya pengembangan rendah."],
        ["Deployment", "Docker Compose, Nginx, GitHub Actions", "Mudah direplikasi pada VPS dengan pengujian otomatis sebelum deploy."],
    ], [2.4, 4.2, 6.6]),
    ("p", "Selain modul agen, basis kode mencerminkan fitur autentikasi JWT, CRUD produk, chat AI dan streaming, order, payment, analytics, inbox, campaigns, workflows, billing, storefront, trust profile, growth links, referrals, knowledge base, QA review, experiments, serta admin dashboard. Lapisan JUALIN OS menyatukannya menjadi operasi toko yang otonom dan terkendali."),
    ("bullets", [
        "Keamanan: validasi keamanan produksi, rate limiting, CORS, request logging, JWT, sanitasi upload gambar, validasi signature webhook, dan filter seller_id pada route penting.",
        "Kualitas dan keamanan AI: tujuh guardrail, kebijakan untrusted data anti prompt-injection, mesin negosiasi deterministik aman-margin, dan trace untuk evaluasi.",
        "Keandalan: idempotency key, background job dengan retry, savepoint transaksi pada lapisan agen, dan endpoint health serta readiness.",
        "Tata kelola otonomi: agent_policies untuk kebijakan per penjual, agent_approvals untuk persetujuan manusia, dan agent_runs untuk activity feed yang dapat diaudit.",
    ]),
    ("table", [
        ["Verifikasi", "Hasil"],
        ["compileall backend", "Berhasil; seluruh modul backend termasuk services agent_os lolos kompilasi."],
        ["Validasi mesin negosiasi", "Nol pelanggaran lantai margin pada seluruh skenario dan ronde uji."],
        ["CI GitHub Actions", "Dikonfigurasi menjalankan pytest, audit dependensi, lint, dan build frontend pada setiap perubahan."],
    ], [5.0, 8.2]),
])

add_section("8. Screenshot Mockup Interface Perangkat Lunak", [
    ("p", "Antarmuka dirancang familiar untuk UMKM sekaligus menampilkan kendali atas tim agen. Customer berinteraksi melalui chat seperti aplikasi percakapan, termasuk saat tawar-menawar yang dijawab agen Juru Tawar secara aman-margin. Penjual memantau melalui dashboard yang fokus pada uang dan order, sementara Pusat Komando AI Crew menampilkan status tiap agen, activity feed, antrean persetujuan, dan laporan harian sehingga penjual selalu tahu apa yang dilakukan AI."),
    ("figure", str(asset_paths[2]), "Gambar 3. Mockup halaman chat customer dengan layanan agen."),
    ("figure", str(asset_paths[3]), "Gambar 4. Mockup dashboard penjual."),
    ("figure", str(asset_paths[4]), "Gambar 5. Mockup quick start onboarding penjual."),
])

add_section("9. Dokumentasi Cara Penggunaan Perangkat Lunak", [
    ("p", "Penggunaan dibagi menjadi dua alur, yaitu penjual sebagai pemilik toko dan customer sebagai pembeli. Dokumentasi ini juga menjadi dasar video demo maksimal tiga menit pada babak final."),
    ("table", [
        ["Langkah seller", "Deskripsi"],
        ["1. Registrasi dan setup", "Penjual membuat akun toko, mengisi katalog dan modal produk untuk batas margin, lalu memilih gaya AI."],
        ["2. Atur kebijakan agen", "Penjual menyetel kebijakan di Pusat Komando AI Crew: diskon maksimum, margin minimum, dan ambang persetujuan."],
        ["3. Aktifkan AI Crew", "Tim agen mulai melayani chat, menawar, menjaga stok, dan menagih secara otonom."],
        ["4. Pantau activity feed", "Penjual melihat tindakan tiap agen secara langsung beserta alasan keputusan."],
        ["5. Setujui keputusan berisiko", "Penjual menyetujui atau menolak diskon besar dan aksi sensitif dari antrean persetujuan."],
        ["6. Baca Laporan Harian", "Penjual menerima ringkasan omzet, pembayaran tertunda, stok menipis, dan saran tindakan."],
    ], [3.2, 10.0]),
    ("table", [
        ["Langkah customer", "Deskripsi"],
        ["1. Buka link chat", "Customer membuka halaman chat publik toko tanpa login."],
        ["2. Tanya dan menawar", "Customer menanyakan produk atau menawar; agen Juru Tawar menanggapi dengan penawaran aman-margin."],
        ["3. Konfirmasi order", "Customer menyetujui detail item, jumlah, dan data pengiriman."],
        ["4. Bayar", "Customer membuka link pembayaran resmi dan memilih QRIS atau VA bila gateway aktif."],
        ["5. Tindak lanjut", "Bila belum bayar, agen Marketing menindaklanjuti sesuai aturan yang diaudit."],
    ], [3.2, 10.0]),
    ("p", "Skenario demo yang disarankan: penjual mengatur kebijakan, customer menawar sebuah produk di bawah batas margin, agen Juru Tawar menawar balik pada harga aman, customer setuju, agen Pramuniaga membuat order dan link pembayaran, lalu penjual melihat jejak ketiga agen pada activity feed dan menerima laporan harian."),
])

add_section("10. Ucapan Terima Kasih", [
    ("p", "Tim Digiboom mengucapkan terima kasih kepada Program Studi Sistem Informasi UNUSA, Himpunan Mahasiswa Sistem Informasi UNUSA, dosen pembimbing, panitia GEMASTIK DIGINEXS 2026, serta rekan-rekan mahasiswa dan pelaku UMKM yang bersedia memberi masukan terhadap pengembangan JUALIN.AI."),
    ("p", "Apresiasi juga diberikan kepada komunitas open-source yang menyediakan fondasi teknologi seperti FastAPI, Next.js, PostgreSQL, pgvector, Redis, dan Docker, sehingga mahasiswa dapat membangun prototipe software berkualitas tinggi dengan biaya pengembangan yang tetap masuk akal."),
])

add_section("11. Daftar Pustaka", [("refs", refs)])

add_section("12. Lampiran", [
    ("p", "Lampiran A. Instrumen validasi 30 responden asli. Instrumen ini disiapkan untuk GForm dan tidak boleh diisi sebagai data fiktif. Rekap akhir harus berasal dari jawaban responden yang benar-benar mengisi."),
    ("table", [
        ["No", "Pertanyaan", "Tipe jawaban"],
        ["1", "Profil responden: pemilik UMKM, reseller, admin online shop, calon pembeli aktif, atau lainnya", "Pilihan tunggal"],
        ["2", "Kanal yang paling sering digunakan untuk jual-beli online", "Pilihan ganda"],
        ["3", "Kendala terbesar saat jualan/berbelanja lewat chat", "Pilihan ganda"],
        ["4", "Seberapa sering chat customer terlambat dibalas atau tidak terjawab", "Skala 1-5"],
        ["5", "Seberapa penting asisten chat 24/7 yang paham katalog produk", "Skala 1-5"],
        ["6", "Fitur JUALIN.AI yang paling dibutuhkan", "Pilihan ganda"],
        ["7", "Seberapa besar minat mencoba JUALIN.AI setelah melihat deskripsi/demo", "Skala 1-5"],
        ["8", "Harga bulanan yang masih dianggap wajar untuk UMKM mikro", "Pilihan tunggal"],
        ["9", "Kekhawatiran terbesar terhadap penggunaan AI dalam chat penjualan", "Pilihan ganda"],
        ["10", "Saran atau kebutuhan tambahan", "Jawaban singkat"],
    ], [1.0, 8.2, 4.0]),
    ("p", "Lampiran B. Rencana rekap validasi. File Excel terpisah sudah disiapkan dengan 30 baris responden kosong, kolom jawaban, dan formula ringkasan. Tim hanya perlu menyalin hasil GForm asli ke sheet Respon_Asli."),
    ("table", [
        ["Indikator", "Target interpretasi"],
        ["Minat mencoba", "Rata-rata skor 4 atau lebih menunjukkan produk layak diprioritaskan untuk uji coba UMKM."],
        ["Kebutuhan asisten chat", "Mayoritas skor 4-5 menunjukkan masalah slow response relevan."],
        ["Fitur prioritas", "Fitur dengan pilihan tertinggi menjadi fokus demo: chat AI, katalog, order, payment, atau dashboard."],
        ["Kekhawatiran", "Kekhawatiran dominan diterjemahkan menjadi guardrails, manual takeover, dan edukasi penggunaan."],
    ], [4.0, 9.2]),
    ("p", "Lampiran C. Hasil uji similaritas. Sesuai guidebook, hasil uji similaritas maksimal 20% perlu dilampirkan oleh tim setelah pengecekan menggunakan Turnitin, iThenticate, atau alat sejenis. Dokumen ini belum menyertakan hasil tersebut karena akses alat similaritas berada pada tim/panitia/kampus."),
])


def build_workbook():
    wb = Workbook()
    ws = wb.active
    ws.title = "Pertanyaan GForm"
    ws.append(["No", "Pertanyaan", "Tipe", "Opsi yang disarankan"])
    questions = [
        (1, "Profil responden", "Pilihan tunggal", "Pemilik UMKM; Reseller/admin online shop; Calon pembeli aktif; Mahasiswa yang sering belanja online; Lainnya"),
        (2, "Kanal jual-beli online yang paling sering digunakan", "Pilihan ganda", "WhatsApp; Instagram; TikTok Shop; Marketplace; Facebook; Website toko"),
        (3, "Kendala terbesar dalam transaksi lewat chat", "Pilihan ganda", "Lambat dibalas; Stok/harga tidak jelas; Sulit checkout; Lupa follow-up bayar; Tidak percaya toko; Lainnya"),
        (4, "Frekuensi chat terlambat dibalas/tidak terjawab", "Skala 1-5", "1 Tidak pernah - 5 Sangat sering"),
        (5, "Pentingnya asisten chat 24/7 yang paham katalog", "Skala 1-5", "1 Tidak penting - 5 Sangat penting"),
        (6, "Fitur paling dibutuhkan", "Pilihan ganda", "Chat AI; Katalog produk; Auto order; Payment link; Follow-up; Dashboard analytics; Manual takeover"),
        (7, "Minat mencoba JUALIN.AI", "Skala 1-5", "1 Tidak berminat - 5 Sangat berminat"),
        (8, "Harga bulanan yang wajar", "Pilihan tunggal", "Gratis; < Rp50.000; Rp50.000-100.000; Rp100.000-300.000; > Rp300.000"),
        (9, "Kekhawatiran terbesar", "Pilihan ganda", "AI salah jawab; Data produk bocor; Biaya; Sulit dipakai; Tidak cocok dengan gaya bahasa toko; Lainnya"),
        (10, "Saran/kebutuhan tambahan", "Jawaban singkat", "Tuliskan bebas"),
    ]
    for row in questions:
        ws.append(row)
    ws2 = wb.create_sheet("Respon_Asli")
    cols = ["ID", "Profil", "Kanal", "Kendala", "Frekuensi slow response (1-5)", "Pentingnya AI chat (1-5)", "Fitur prioritas", "Minat mencoba (1-5)", "Harga wajar", "Kekhawatiran", "Saran"]
    ws2.append(cols)
    for i in range(1, 31):
        ws2.append([f"R{i:02d}"] + [""] * (len(cols) - 1))
    ws3 = wb.create_sheet("Ringkasan")
    for row in [
        ["Metrik", "Formula/hasil"],
        ["Jumlah responden terisi", "=COUNTA(Respon_Asli!B2:B31)"],
        ["Rata-rata frekuensi slow response", "=AVERAGEIF(Respon_Asli!E2:E31,\">0\")"],
        ["Rata-rata pentingnya AI chat", "=AVERAGEIF(Respon_Asli!F2:F31,\">0\")"],
        ["Rata-rata minat mencoba", "=AVERAGEIF(Respon_Asli!H2:H31,\">0\")"],
        ["Catatan", "Isi hanya dari respons GForm asli. Jangan memakai data simulasi sebagai bukti proposal."],
    ]:
        ws3.append(row)
    for sheet in [ws, ws2, ws3]:
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="16A34A")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = Border(
                    left=Side(style="thin", color="D9E2EC"),
                    right=Side(style="thin", color="D9E2EC"),
                    top=Side(style="thin", color="D9E2EC"),
                    bottom=Side(style="thin", color="D9E2EC"),
                )
        for idx, width in enumerate([10, 28, 28, 42, 20, 20, 30, 18, 24, 28, 36][:sheet.max_column], start=1):
            sheet.column_dimensions[get_column_letter(idx)].width = width
    ws2.freeze_panes = "A2"
    wb.save(XLSX_PATH)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(10.5)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def docx_add_table(doc, rows, widths_cm):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.cell(ri, ci)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cell, value, bold=(ri == 0), color=("FFFFFF" if ri == 0 else None))
            if ri == 0:
                set_cell_shading(cell, "16A34A")
            elif ri % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
            if ci < len(widths_cm):
                cell.width = Cm(widths_cm[ci])
    doc.add_paragraph()


def docx_add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(12)


def docx_add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def centered(doc, text, size=12, bold=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(size)
    r.bold = bold


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(3)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    for st_name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        st = styles[st_name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.line_spacing = 1.5
    for _ in range(5):
        doc.add_paragraph()
    centered(doc, "PROPOSAL PENGEMBANGAN PERANGKAT LUNAK", 14, True, 12)
    centered(doc, "JUALIN.AI", 18, True, 4)
    centered(doc, "AI Sales Assistant Berbasis Katalog untuk Otomasi Layanan Chat UMKM Mikro", 13, True, 24)
    centered(doc, "Divisi III: Pengembangan Perangkat Lunak (Software Development)", 12, False, 4)
    centered(doc, "GEMASTIK DIGINEXS 2026", 12, False, 24)
    centered(doc, "Tim Digiboom", 12, True, 8)
    centered(doc, "Ketua: [Nama Ketua] / [NIM]", 12, False, 2)
    centered(doc, "Anggota 1: [Nama Anggota 1] / [NIM]", 12, False, 2)
    centered(doc, "Anggota 2: [Nama Anggota 2] / [NIM]", 12, False, 24)
    centered(doc, "Program Studi Sistem Informasi", 12, False, 2)
    centered(doc, "Universitas Nahdlatul Ulama Surabaya", 12, False, 2)
    centered(doc, TODAY, 12, False, 2)
    doc.add_page_break()
    for title, blocks in sections:
        doc.add_heading(title, level=1)
        for block in blocks:
            kind = block[0]
            if kind == "p":
                docx_add_para(doc, block[1])
            elif kind == "bullets":
                docx_add_bullets(doc, block[1])
            elif kind == "table":
                docx_add_table(doc, block[1], block[2])
            elif kind == "figure":
                doc.add_picture(block[1], width=Inches(5.8))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(8)
                r = cap.add_run(block[2])
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
                r.italic = True
            elif kind == "refs":
                for i, ref in enumerate(block[1], start=1):
                    docx_add_para(doc, f"[{i}] {ref}")
        doc.add_paragraph()
    doc.save(DOCX_PATH)


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BodyJ", parent=styles["Normal"], fontName="TNR", fontSize=12, leading=18, alignment=TA_JUSTIFY, spaceAfter=7))
    styles.add(ParagraphStyle(name="BodyL", parent=styles["Normal"], fontName="TNR", fontSize=12, leading=18, alignment=TA_LEFT, spaceAfter=7))
    styles.add(ParagraphStyle(name="TitleCenter", parent=styles["Title"], fontName="TNR-Bold", fontSize=18, leading=24, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="SubCenter", parent=styles["Normal"], fontName="TNR", fontSize=12, leading=18, alignment=TA_CENTER, spaceAfter=6))
    styles.add(ParagraphStyle(name="H1x", parent=styles["Heading1"], fontName="TNR-Bold", fontSize=14, leading=21, alignment=TA_LEFT, spaceBefore=12, spaceAfter=8))
    styles.add(ParagraphStyle(name="Caption", parent=styles["Normal"], fontName="TNR-Italic", fontSize=10, leading=13, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="Cell", parent=styles["Normal"], fontName="TNR", fontSize=9.6, leading=12, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="CellHead", parent=styles["Normal"], fontName="TNR-Bold", fontSize=9.6, leading=12, alignment=TA_CENTER, textColor=colors.white))

    def para(text, style="BodyJ"):
        return Paragraph(html.escape(str(text)), styles[style])

    def make_table(rows, widths_cm):
        data = []
        for ri, row in enumerate(rows):
            data.append([Paragraph(html.escape(str(c)), styles["CellHead" if ri == 0 else "Cell"]) for c in row])
        t = Table(data, colWidths=[w * cm for w in widths_cm], hAlign="CENTER", repeatRows=1)
        ts = TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#16A34A")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.45, HexColor("#CBD5E1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ])
        for r in range(1, len(rows)):
            if r % 2 == 0:
                ts.add("BACKGROUND", (0, r), (-1, r), HexColor("#F8FAFC"))
        t.setStyle(ts)
        return t

    def make_bullets(items):
        return ListFlowable(
            [ListItem(para(i, "BodyL"), leftIndent=12) for i in items],
            bulletType="bullet",
            leftIndent=22,
            bulletFontName="TNR",
            bulletFontSize=10,
        )

    def add_figure(story, path, caption):
        img = RLImage(path)
        max_w = 13.0 * cm
        ratio = max_w / img.imageWidth
        img.drawWidth = max_w
        img.drawHeight = img.imageHeight * ratio
        story.append(KeepTogether([img, para(caption, "Caption")]))

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("TNR", 9)
        canvas.setFillColor(colors.grey)
        canvas.drawString(4 * cm, 1.5 * cm, "JUALIN.AI - Proposal Software Development")
        canvas.drawRightString(A4[0] - 3 * cm, 1.5 * cm, f"Halaman {doc.page}")
        canvas.restoreState()

    story = []
    story.append(Spacer(1, 4.5 * cm))
    story.append(para("PROPOSAL PENGEMBANGAN PERANGKAT LUNAK", "SubCenter"))
    story.append(Paragraph("JUALIN.AI", styles["TitleCenter"]))
    story.append(para("Sistem Operasi Toko Otonom (JUALIN OS) untuk UMKM Mikro", "SubCenter"))
    story.append(Spacer(1, 0.8 * cm))
    story.append(para("Divisi III: Pengembangan Perangkat Lunak (Software Development)", "SubCenter"))
    story.append(para("GEMASTIK DIGINEXS 2026", "SubCenter"))
    story.append(Spacer(1, 1.2 * cm))
    story.append(para("Tim Digiboom", "SubCenter"))
    story.append(para("Ketua: [Nama Ketua] / [NIM]", "SubCenter"))
    story.append(para("Anggota 1: [Nama Anggota 1] / [NIM]", "SubCenter"))
    story.append(para("Anggota 2: [Nama Anggota 2] / [NIM]", "SubCenter"))
    story.append(Spacer(1, 1.2 * cm))
    story.append(para("Program Studi Sistem Informasi", "SubCenter"))
    story.append(para("Universitas Nahdlatul Ulama Surabaya", "SubCenter"))
    story.append(para(TODAY, "SubCenter"))
    story.append(PageBreak())
    for title, blocks in sections:
        story.append(Paragraph(html.escape(title), styles["H1x"]))
        for block in blocks:
            kind = block[0]
            if kind == "p":
                story.append(para(block[1]))
            elif kind == "bullets":
                story.append(make_bullets(block[1]))
            elif kind == "table":
                story.append(make_table(block[1], block[2]))
                story.append(Spacer(1, 0.25 * cm))
            elif kind == "figure":
                add_figure(story, block[1], block[2])
            elif kind == "refs":
                for i, ref in enumerate(block[1], start=1):
                    story.append(para(f"[{i}] {ref}", "BodyL"))
        story.append(Spacer(1, 0.2 * cm))
    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, leftMargin=4 * cm, rightMargin=3 * cm, topMargin=3 * cm, bottomMargin=3 * cm)
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def build_notes():
    NOTES_PATH.write_text(
        "Catatan finalisasi submission JUALIN.AI\n"
        "1. Isi nama ketua, anggota, NIM, dan program studi pada cover DOCX/PDF.\n"
        "2. Jalankan GForm dengan pertanyaan pada Template_Kuesioner_Validasi_JUALIN_AI.xlsx. Gunakan jawaban asli, bukan data dummy.\n"
        "3. Setelah respons masuk, salin ke sheet Respon_Asli dan update bagian Lampiran/Ringkasan jika ingin mencantumkan hasil.\n"
        "4. Jalankan uji similaritas; lampirkan hasil maksimal 20% sesuai guidebook.\n"
        "5. File PDF utama untuk upload: SoftwareDevelopment_Digiboom.pdf.\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    build_workbook()
    build_docx()
    build_pdf()
    build_notes()
    from pypdf import PdfReader
    reader = PdfReader(str(PDF_PATH))
    print("CREATED", DOCX_PATH)
    print("CREATED", PDF_PATH, "pages=", len(reader.pages))
    print("CREATED", XLSX_PATH)
    print("CREATED", NOTES_PATH)
    print("ASSETS", [p.name for p in asset_paths])
